import json
import boto3
from docx import Document
import io

MAX_WORDS = 80
MAX_TOKENS = int(MAX_WORDS * 1.5)
PROMPT_EHC_JSON_S3_PATH = 's3://artifacts-gen-ai-eu-west-2/prompt_engineering_json/EHC_Plan_prompts.json'

possible_topics = [
    "Child_Young_Person_History",
    "Views_Interests_Strengths_Aspirations",
    "Special_Educational_Needs_Provision",
    "Health_Care_Needs_Related_to_SEN",
    "Social_Care_Needs"
]


def read_json_from_s3(bucket, key):
    s3 = boto3.client('s3')
    response = s3.get_object(Bucket=bucket, Key=key)
    content = response['Body'].read().decode('utf-8')
    return json.loads(content)


def read_docx_from_s3(bucket, key):
    s3 = boto3.client('s3')
    response = s3.get_object(Bucket=bucket, Key=key)
    doc = Document(io.BytesIO(response['Body'].read()))
    return '\n'.join([paragraph.text for paragraph in doc.paragraphs])


def call_claude_sonnet(prompt, max_tokens=MAX_TOKENS):
    bedrock_runtime = boto3.client(
        service_name='bedrock-runtime',
        region_name='eu-west-2'
    )

    prompt_config = {
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": max_tokens * 3,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                ],
            }
        ],
    }

    body = json.dumps(prompt_config)

    modelId = "anthropic.claude-3-sonnet-20240229-v1:0"
    accept = "application/json"
    contentType = "application/json"

    response = bedrock_runtime.invoke_model(
        body=body, modelId=modelId, accept=accept, contentType=contentType
    )

    response_body = json.loads(response.get("body").read())
    result = response_body.get("content")[0].get("text")
    return result


def generate_section(topic, input_form, ehc_plan_prompts):
    if topic not in possible_topics:
        print(f"Error: {topic} not in {possible_topics}")
        return ""
    prompt = """Below is a structured form that has been filled out. 
    Use this data to perform the task which will be assigned to you below that form.
    Do not say you are responding to a prompt or task. Limit your response to {max_words} sentences.
        <structured form data>
            {structured_form}
        </structured form data>
        <task>
            {task_prompt}
        </task>
    """
    full_prompt = prompt.format(
        structured_form=input_form,
        max_words=MAX_WORDS,
        task_prompt=ehc_plan_prompts['EHC_Plan_Prompts'][topic]['Prompt']
    )

    response = call_claude_sonnet(full_prompt)
    if len(response.split()) > MAX_WORDS:
        print("Response too long, shrinking")
        shrinking_prompt = """Reduce the number of words in the below text to be at most {max_words} words,
        and at least {min_words} words.
        IMPORTANT: do not change the meaning of the text or remove any information:
        <TEXT>
        {text}
        </TEXT>"""
        full_prompt = shrinking_prompt.format(
            text=response,
            max_words=MAX_WORDS,
            min_words=int(MAX_WORDS * 0.8)
        )
        response = call_claude_sonnet(full_prompt)
    return response


def lambda_handler(event, context):
    print(event)
    # Extract S3 bucket and key from the event
    source_bucket = event['Records'][0]['s3']['bucket']['name']
    source_key = event['Records'][0]['s3']['object']['key'].replace('+', ' ')

    # Read the input form from S3
    input_form = read_docx_from_s3(source_bucket, source_key)

    # Parse S3 bucket and key from the PROMPT_EHC_JSON_S3_PATH
    s3_parts = PROMPT_EHC_JSON_S3_PATH.split('/', 3)
    prompt_bucket = s3_parts[2]
    prompt_key = s3_parts[3]

    # Read EHC Plan prompts from S3
    ehc_plan_prompts = read_json_from_s3(prompt_bucket, prompt_key)

    document = Document()

    for topic in possible_topics:
        print(f"topic: {topic}")
        response = generate_section(topic, input_form, ehc_plan_prompts)
        print(response)
        print(len(response.split()))

        document.add_heading(topic.replace("_", " "), level=1)
        document.add_paragraph(response)

    # Save the document to a BytesIO object
    docx_buffer = io.BytesIO()
    document.save(docx_buffer)
    docx_buffer.seek(0)

    # Generate the destination key
    destination_key = f"processed_{source_key.split('/')[-1]}"
    print(f"destination_key {destination_key}")

    # Upload the document to S3
    s3 = boto3.client('s3')
    try:
        s3.put_object(
            Bucket='clean-ehc-report-eu-west-2',
            Key=destination_key,  # Make sure 'Key' is capitalized
            Body=docx_buffer.getvalue()
        )
        print(f"Successfully uploaded to S3: {destination_key}")
    except Exception as e:
        print(f"Error uploading to S3: {str(e)}")
        raise

    return {
        'statusCode': 200,
        'body': json.dumps('EHCP document generated and saved to S3 successfully')
    }