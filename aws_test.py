# For printing and system operations
import pprint
# For printing and system operations
import sys
# AWS SDK for Python to interact with AWS services
import boto3
# To handle JSON data
import json

# To create Word documents
from docx import Document

# Limits for the number of words and tokens in the summarised output
MAX_WORDS = 80
MAX_TOKENS = int(MAX_WORDS*1.5)
# Reads prompts from a JSON file to be used for generating summaries
EHC_Plan_prompts = eval(open("EHC_Plan_prompts.json", "r").read())
# pprint.pprint(EHC_Plan_prompts['EHC_Plan_Prompts'])

# Lists topics related to the EHCP sections
possible_topics = ["Child_Young_Person_History",
                            "Views_Interests_Strengths_Aspirations",
                            "Special_Educational_Needs_Provision",
                            "Health_Care_Needs_Related_to_SEN",
                            "Social_Care_Needs"]

# This function calls the Anthropic Claude 3 Sonnet model to generate text based on a given prompt
def call_claude_sonnet(prompt, max_tokens=MAX_TOKENS):
    bedrock_runtime = boto3.client(
        service_name='bedrock-runtime',
        region_name='eu-west-2'
    )

    prompt_config = {
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": max_tokens*3,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                ],
            }
        ],
    }

    body = json.dumps(prompt_config)

    modelId = "anthropic.claude-3-sonnet-20240229-v1:0"
    accept = "application/json"
    contentType = "application/json"

    response = bedrock_runtime.invoke_model(
        body=body, modelId=modelId, accept=accept, contentType=contentType
    )

    response_body = json.loads(response.get("body").read())
    result = response_body.get("content")[0].get("text")
    return result




#This function generates a summarised section based on a given topic and input form. It first creates a prompt using the input form and topic, then calls the call_claude_haiku function to get the generated text
def generate_section(topic, input_form):
    if topic not in possible_topics:
        print(f"Error: {topic} not in {possible_topics}")
        return ""
    prompt = """Below is a structured form that has been filled out. 
    Use this data to perform the task which will be assigned to you below that form.
    Do not say you are responding to a prompt or task. Limit your response to {max_words} sentences.
        <structured form data>
            {structured_form}
        </structured form data>
        <task>
            {task_prompt}
        </task>
    """
    full_prompt = prompt.format(structured_form=input_form,
                                max_words = MAX_WORDS,
                                task_prompt=EHC_Plan_prompts['EHC_Plan_Prompts'][topic]['Prompt'])

    response = call_claude_sonnet(full_prompt)
    if len(response.split()) > MAX_WORDS:
        print("Response too long, shrinking")
        shrinking_prompt = """Reduce the number of words in the below text to be at most {max_words} words,
        and at least {min_words} words.
        IMPORTANT: do not change the meaning of the text or remove any information:
        <TEXT>
        {text}
        </TEXT>"""
        full_prompt = shrinking_prompt.format(text=response, max_words=MAX_WORDS,min_words = int(MAX_WORDS*0.8))
        response = call_claude_sonnet(full_prompt)
    return response

def find_num_words():
    for topic in possible_topics:
        print(topic, len(EHC_Plan_prompts['EHC_Plan_Prompts'][topic]['Example_Output'].split()))


# Example usage
example_input_form = open("ExampleInput1.txt", "r", encoding='latin-1').read()
document = Document()

for topic in possible_topics:
    print(f"topic: {topic}")
    response = generate_section(topic, example_input_form)
    print(response)
    print(len(response.split()))

    document.add_heading(topic.replace("_", " "), level=1)
    document.add_paragraph(response)

# Save the document
document.save("Generated_EHCP.docx")